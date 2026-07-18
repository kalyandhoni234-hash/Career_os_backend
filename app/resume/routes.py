import json
from datetime import datetime, timezone
from io import BytesIO
from flask import Blueprint, request, jsonify, make_response
from flask_login import login_required, current_user
from app.extensions import db, limiter
from app.resume.models import Resume, ResumeVersion
from app.resume.ats import score_resume

resume_bp = Blueprint("resume", __name__)


def _normalize_list_field(items, array_fields):
    """Normalize JSON list entries so specified fields are always arrays."""
    if not isinstance(items, list):
        return []
    out = []
    for item in items:
        if not isinstance(item, dict):
            continue
        entry = dict(item)
        for key in array_fields:
            val = entry.get(key, [])
            if isinstance(val, str):
                entry[key] = [s.strip() for s in val.split("\n") if s.strip()] or [""]
            elif not isinstance(val, list):
                entry[key] = [str(val)] if val else []
            elif key not in entry:
                entry[key] = []
        out.append(entry)
    return out


def _serialize(resume):
    return {
        "id": resume.id,
        "full_name": resume.full_name,
        "email": resume.email,
        "phone": resume.phone,
        "location": resume.location,
        "summary": resume.summary,
        "title": resume.title,
        "website": resume.website,
        "linkedin": resume.linkedin,
        "github": resume.github,
        "portfolio": resume.portfolio,
        "experience": resume.experience,
        "education": resume.education,
        "projects": resume.projects,
        "skills": resume.skills,
        "certificates": resume.certificates,
        "achievements": resume.achievements,
        "languages": resume.languages,
        "publications": resume.publications,
        "tone": resume.tone or "professional",
        "target_job_description": resume.target_job_description,
        "created_at": resume.created_at.isoformat() if resume.created_at else None,
        "updated_at": resume.updated_at.isoformat() if resume.updated_at else None,
    }


def _snapshot(resume):
    return {
        "full_name": resume.full_name,
        "email": resume.email,
        "phone": resume.phone,
        "location": resume.location,
        "summary": resume.summary,
        "title": resume.title,
        "website": resume.website,
        "linkedin": resume.linkedin,
        "github": resume.github,
        "portfolio": resume.portfolio,
        "experience": _normalize_list_field(
            resume.experience, ["bullets", "technologies"]
        ),
        "education": resume.education,
        "projects": _normalize_list_field(resume.projects, ["technologies"]),
        "skills": resume.skills,
        "certificates": resume.certificates,
        "achievements": resume.achievements,
        "languages": resume.languages,
        "publications": resume.publications,
        "tone": resume.tone or "professional",
    }


@resume_bp.route("/ping")
def ping():
    return {"blueprint": "resume", "status": "alive"}


@resume_bp.route("", methods=["GET"])
@login_required
def get_resume():
    from app.resume.profile_bridge import get_merged_resume as _get_merged

    merged = _get_merged(current_user.id)
    if merged is None:
        return jsonify({"resume": None}), 200
    return jsonify({"resume": merged}), 200


@resume_bp.route("", methods=["POST", "PUT"])
@login_required
def upsert_resume():
    data = request.get_json(silent=True) or {}
    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        resume = Resume(user_id=current_user.id)
        db.session.add(resume)

    for field in [
        "full_name",
        "email",
        "phone",
        "location",
        "summary",
        "title",
        "website",
        "linkedin",
        "github",
        "portfolio",
        "experience",
        "education",
        "projects",
        "skills",
        "certificates",
        "achievements",
        "languages",
        "publications",
        "tone",
    ]:
        if field in data:
            value = data[field]
            if field == "experience":
                value = _normalize_list_field(value, ["bullets", "technologies"])
            elif field == "projects":
                value = _normalize_list_field(value, ["technologies"])
            setattr(resume, field, value)

    db.session.commit()

    # Propagate to canonical profile tables (single-source-of-truth)
    from app.resume.profile_bridge import save_resume_to_canonical
    save_resume_to_canonical(current_user.id, data)

    # Create a version snapshot
    version_name = data.get("version_name") or f"v{resume.versions.count() + 1}"
    version = ResumeVersion(
        resume_id=resume.id,
        version_name=version_name,
        target_role=data.get("target_role", data.get("title", "")),
        source="manual",
        snapshot=_snapshot(resume),
    )
    db.session.add(version)
    db.session.commit()

    from app.core.integration import on_resume_changed
    on_resume_changed(current_user.id)

    return jsonify({"message": "Resume saved successfully", "id": resume.id}), 200


def _snapshot_to_html(snapshot):
    """Convert a resume snapshot dict into a printable HTML string."""
    def esc(text):
        if not text:
            return ""
        return (
            str(text)
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#39;")
        )

    items_html = ""

    if snapshot.get("summary"):
        items_html += f'<div class="section"><h2>Professional Summary</h2><p>{esc(snapshot["summary"])}</p></div>'

    exp_entries = ""
    for exp in snapshot.get("experience", []) or []:
        role = esc(exp.get("role", ""))
        company = esc(exp.get("company", ""))
        start = esc(exp.get("start", ""))
        end = esc(exp.get("end", ""))
        bullets = exp.get("bullets", [])
        if isinstance(bullets, str):
            bullets = bullets.split("\n")
        bullet_items = "".join(f"<li>{esc(b)}</li>" for b in bullets if b.strip())
        tech = exp.get("technologies", "")
        tech_str = (
            f'<p class="tech"><strong>Technologies:</strong> {esc(", ".join(tech) if isinstance(tech, list) else tech)}</p>'
            if tech
            else ""
        )
        exp_entries += f'<div class="entry"><div class="entry-header"><strong>{role}</strong> at {company}</div><div class="date">{start} - {end}</div><ul>{bullet_items}</ul>{tech_str}</div>'
    if exp_entries:
        items_html += f'<div class="section"><h2>Experience</h2>{exp_entries}</div>'

    edu_entries = ""
    for edu in snapshot.get("education", []) or []:
        school = esc(edu.get("school", ""))
        degree = esc(edu.get("degree", ""))
        field = esc(edu.get("field", ""))
        start = esc(edu.get("start", ""))
        end = edu.get("end", "")
        gpa = esc(edu.get("gpa", ""))
        gpa_str = f'<span class="gpa">GPA: {gpa}</span>' if gpa else ""
        edu_entries += f'<div class="entry"><div class="entry-header"><strong>{degree}</strong> in {field}</div><div class="entry-sub">{school}</div><div class="date">{start} - {end} {gpa_str}</div></div>'
    if edu_entries:
        items_html += f'<div class="section"><h2>Education</h2>{edu_entries}</div>'

    proj_entries = ""
    for proj in snapshot.get("projects", []) or []:
        name = esc(proj.get("name", ""))
        desc = esc(proj.get("description", ""))
        tech = proj.get("technologies", "")
        tech_str = (
            f'<p class="tech"><strong>Technologies:</strong> {esc(", ".join(tech) if isinstance(tech, list) else tech)}</p>'
            if tech
            else ""
        )
        url = proj.get("url", "")
        url_str = f'<p><a href="{esc(url)}">{esc(url)}</a></p>' if url else ""
        proj_entries += f'<div class="entry"><div class="entry-header"><strong>{name}</strong></div><p>{desc}</p>{tech_str}{url_str}</div>'
    if proj_entries:
        items_html += f'<div class="section"><h2>Projects</h2>{proj_entries}</div>'

    skills = snapshot.get("skills", [])
    if skills:
        if isinstance(skills, list):
            skills_str = esc(", ".join(skills))
        else:
            skills_str = esc(str(skills))
        items_html += f'<div class="section"><h2>Skills</h2><p>{skills_str}</p></div>'

    cert_html = ""
    for cert in snapshot.get("certificates", []) or []:
        cert_html += f'<div class="entry"><strong>{esc(cert.get("name", ""))}</strong> - {esc(cert.get("issuer", ""))} ({esc(cert.get("date", ""))})</div>'
    if cert_html:
        items_html += f'<div class="section"><h2>Certificates</h2>{cert_html}</div>'

    langs = snapshot.get("languages", [])
    if langs:
        lang_str = esc(
            ", ".join(
                f"{lang.get('name', '')} ({lang.get('level', '')})"
                if isinstance(lang, dict)
                else str(lang)
                for lang in langs
            )
        )
        items_html += f'<div class="section"><h2>Languages</h2><p>{lang_str}</p></div>'

    name = esc(snapshot.get("full_name", "Resume"))
    contact = ", ".join(
        filter(None, [esc(snapshot.get("email", "")), esc(snapshot.get("phone", "")), esc(snapshot.get("location", ""))])
    )

    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
@page {{ margin: 40px; }}
body {{ font-family: 'Segoe UI', Arial, Helvetica, sans-serif; font-size: 12px; color: #1a1a1a; line-height: 1.5; margin: 0; padding: 0; }}
h1 {{ font-size: 22px; margin: 0 0 4px 0; color: #111; }}
.contact {{ color: #555; font-size: 11px; margin-bottom: 16px; }}
.section {{ margin-bottom: 16px; page-break-inside: avoid; }}
.section h2 {{ font-size: 13px; text-transform: uppercase; letter-spacing: 1px; color: #333; border-bottom: 1.5px solid #333; padding-bottom: 3px; margin: 0 0 8px 0; }}
.entry {{ margin-bottom: 10px; }}
.entry-header {{ font-size: 12px; margin-bottom: 1px; }}
.entry-sub {{ color: #555; font-size: 11px; }}
.date {{ font-size: 10px; color: #777; margin-bottom: 4px; }}
ul {{ margin: 4px 0; padding-left: 18px; }}
li {{ font-size: 11px; margin-bottom: 2px; }}
.tech {{ font-size: 10px; color: #555; margin: 2px 0; }}
.gpa {{ color: #555; }}
a {{ color: #2563eb; text-decoration: none; }}
</style></head><body>
<h1>{name}</h1>
<div class="contact">{contact}</div>
{items_html}
</body></html>"""


def _snapshot_to_docx(snapshot):
    """Convert a resume snapshot dict into a python-docx Document."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    def add_section_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        run.font.name = "Calibri"
        p_fmt = p.paragraph_format
        p_fmt.space_before = Pt(8)
        p_fmt.space_after = Pt(3)
        p_fmt.border_bottom = True

    name_str = snapshot.get("full_name", "Resume")
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_name.add_run(name_str)
    run.bold = True
    run.font.size = Pt(20)

    contact_str = ", ".join(filter(None, [snapshot.get("email", ""), snapshot.get("phone", ""), snapshot.get("location", "")]))
    if contact_str:
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_contact.add_run(contact_str)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if snapshot.get("summary"):
        add_section_heading("Professional Summary")
        doc.add_paragraph(snapshot["summary"])

    if snapshot.get("experience"):
        add_section_heading("Experience")
    for exp in snapshot.get("experience", []) or []:
        role = exp.get("role", "")
        company = exp.get("company", "")
        start = exp.get("start", "")
        end = exp.get("end", "")
        header = f"{role} at {company}" if company else role
        p = doc.add_paragraph()
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10.5)
        if start or end:
            p.add_run(f"    {start} - {end}")
        bullets = exp.get("bullets", [])
        if isinstance(bullets, str):
            bullets = bullets.split("\n")
        for b in bullets:
            if b.strip():
                doc.add_paragraph(b.strip(), style="List Bullet")
        tech = exp.get("technologies", "")
        if tech:
            t = ", ".join(tech) if isinstance(tech, list) else tech
            p_tech = doc.add_paragraph(f"Technologies: {t}")
            p_tech.paragraph_format.space_before = Pt(0)

    if snapshot.get("education"):
        add_section_heading("Education")
    for edu in snapshot.get("education", []) or []:
        school = edu.get("school", "")
        degree = edu.get("degree", "")
        field = edu.get("field", "")
        start = edu.get("start", "")
        end = edu.get("end", "")
        gpa = edu.get("gpa", "")
        line = f"{degree} in {field}" if field else degree
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.bold = True
        sub = f"{school}    {start} - {end}"
        if gpa:
            sub += f"    GPA: {gpa}"
        p_sub = doc.add_paragraph(sub)
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(2)
        for r in p_sub.runs:
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if snapshot.get("projects"):
        add_section_heading("Projects")
    for proj in snapshot.get("projects", []) or []:
        name = proj.get("name", "")
        desc = proj.get("description", "")
        tech = proj.get("technologies", "")
        url = proj.get("url", "")
        p = doc.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        if desc:
            doc.add_paragraph(desc)
        if tech:
            t = ", ".join(tech) if isinstance(tech, list) else tech
            doc.add_paragraph(f"Technologies: {t}")
        if url:
            doc.add_paragraph(url)

    skills = snapshot.get("skills", [])
    if skills:
        add_section_heading("Skills")
        if isinstance(skills, list):
            doc.add_paragraph(", ".join(skills))
        else:
            doc.add_paragraph(str(skills))

    if snapshot.get("certificates"):
        add_section_heading("Certificates")
    for cert in snapshot.get("certificates", []) or []:
        name = cert.get("name", "")
        issuer = cert.get("issuer", "")
        date = cert.get("date", "")
        parts = filter(None, [name, issuer, date])
        doc.add_paragraph(" — ".join(parts))

    if snapshot.get("languages"):
        add_section_heading("Languages")
    langs = snapshot.get("languages", [])
    if langs:
        lang_str = ", ".join(
            f"{lang.get('name', '')} ({lang.get('level', '')})"
            if isinstance(lang, dict)
            else str(lang)
            for lang in langs
        )
        if lang_str:
            doc.add_paragraph(lang_str)

    return doc


@resume_bp.route("/export", methods=["GET"])
@login_required
def export_resume():
    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        return jsonify({"error": "No resume found"}), 404
    html = _snapshot_to_html(_snapshot(resume))
    from app.resume.pdf_engine import html_to_pdf
    pdf_bytes, err = html_to_pdf(html, resume.full_name or "resume")
    if err:
        return jsonify({"error": err}), 503
    response = make_response(pdf_bytes)
    response.headers["Content-Type"] = "application/pdf"
    response.headers["Content-Disposition"] = (
        f"attachment; filename={resume.full_name or 'resume'}.pdf"
    )
    return response


@resume_bp.route("/export/docx", methods=["GET"])
@login_required
def export_resume_docx():
    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        return jsonify({"error": "No resume found"}), 404
    doc = _snapshot_to_docx(_snapshot(resume))
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    response = make_response(bio.read())
    response.headers["Content-Type"] = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response.headers["Content-Disposition"] = (
        f"attachment; filename={resume.full_name or 'resume'}.docx"
    )
    return response


@resume_bp.route("/versions/<int:version_id>/export", methods=["GET"])
@login_required
def export_version_pdf(version_id):
    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        return jsonify({"error": "No resume found"}), 404
    version = ResumeVersion.query.filter_by(id=version_id, resume_id=resume.id).first()
    if not version:
        return jsonify({"error": "Version not found"}), 404
    snapshot = version.snapshot or {}
    html = _snapshot_to_html(snapshot)
    from app.resume.pdf_engine import html_to_pdf
    pdf_bytes, err = html_to_pdf(html, snapshot.get("full_name", "resume") or "resume")
    if err:
        return jsonify({"error": err}), 503
    response = make_response(pdf_bytes)
    response.headers["Content-Type"] = "application/pdf"
    response.headers["Content-Disposition"] = (
        f"attachment; filename={snapshot.get('full_name', 'resume') or 'resume'}.pdf"
    )
    return response


@resume_bp.route("/versions/<int:version_id>/export/docx", methods=["GET"])
@login_required
def export_version_docx(version_id):
    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        return jsonify({"error": "No resume found"}), 404
    version = ResumeVersion.query.filter_by(id=version_id, resume_id=resume.id).first()
    if not version:
        return jsonify({"error": "Version not found"}), 404
    snapshot = version.snapshot or {}
    doc = _snapshot_to_docx(snapshot)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    response = make_response(bio.read())
    response.headers["Content-Type"] = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response.headers["Content-Disposition"] = (
        f"attachment; filename={snapshot.get('full_name', 'resume') or 'resume'}.docx"
    )
    return response


# ── PDF Engine Health ──────────────────────────────────────


@resume_bp.route("/pdf-health", methods=["GET"])
def pdf_health():
    from app.resume.pdf_engine import get_status

    return jsonify(get_status()), 200


# ── AI Resume Review ──────────────────────────────────────


@resume_bp.route("/review", methods=["POST"])
@login_required
@limiter.limit("5 per minute")
def review_resume():
    from app.ai_service import generate_text

    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        return jsonify({"error": "No resume found"}), 404

    experience_lines = []
    for exp in resume.experience or []:
        role = exp.get("role", "")
        company = exp.get("company", "")
        bullets = exp.get("bullets", "")
        if isinstance(bullets, list):
            bullets = "\n".join(f"  - {b}" for b in bullets)
        else:
            bullets = f"  - {bullets}" if bullets else ""
        tech = exp.get("technologies", "")
        if isinstance(tech, list):
            tech = ", ".join(tech)
        experience_lines.append(
            f"- {role} at {company}\n{bullets}\n  Technologies: {tech}"
        )

    project_lines = []
    for proj in resume.projects or []:
        pname = proj.get("name", "")
        pdesc = proj.get("description", "")
        tech = proj.get("technologies", "")
        if isinstance(tech, list):
            tech = ", ".join(tech)
        project_lines.append(f"- {pname}: {pdesc} [{tech}]")

    skills_str = ", ".join(resume.skills or [])
    certs_str = ", ".join(
        f"{c.get('name', '')} ({c.get('issuer', '')})"
        for c in (resume.certificates or [])
    )
    lang_str = ", ".join(
        f"{lang.get('name', '')} ({lang.get('level', '')})"
        if isinstance(lang, dict)
        else str(lang)
        for lang in (resume.languages or [])
    )

    resume_text = f"""
Full Name: {resume.full_name or "N/A"}
Title: {resume.title or "N/A"}
Email: {resume.email or "N/A"}
Phone: {resume.phone or "N/A"}
Location: {resume.location or "N/A"}
Summary: {resume.summary or "N/A"}
Skills: {skills_str}
Certificates: {certs_str or "None"}
Languages: {lang_str or "None"}
Experience:
{chr(10).join(experience_lines) or "None"}
Projects:
{chr(10).join(project_lines) or "None"}
"""

    system_instruction = """You are an expert ATS resume reviewer. Analyze the resume and respond ONLY in valid JSON with this exact structure:
{
  "ats_score": <number 0-100>,
  "strengths": ["...", "..."],
  "weaknesses": ["...", "..."],
  "missing_keywords": ["...", "..."],
  "weak_action_verbs": ["...", "..."],
  "suggestions": ["...", "..."]
}
Do not include any text outside the JSON object."""

    raw_response = generate_text(
        resume_text, model="gemini", system_instruction=system_instruction
    )

    cleaned = raw_response.strip()
    if cleaned.startswith("```"):
        parts = cleaned.split("```")
        cleaned = parts[1] if len(parts) > 1 else cleaned
        if cleaned.startswith("json"):
            cleaned = cleaned[4:]

    try:
        review_data = json.loads(cleaned)
    except json.JSONDecodeError:
        return jsonify(
            {"error": "Failed to parse AI response", "raw": raw_response}
        ), 500

    return jsonify({"review": review_data}), 200


# ── ATS Scoring ───────────────────────────────────────────


@resume_bp.route("/ats-score", methods=["POST"])
@login_required
def ats_score():
    data = request.get_json(silent=True) or {}
    job_description = data.get("job_description", "")

    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        return jsonify({"error": "No resume found"}), 404

    resume.target_job_description = job_description
    db.session.commit()

    result = score_resume(resume, job_description)
    return jsonify(result), 200


# ── Resume Versions ───────────────────────────────────────


@resume_bp.route("/versions", methods=["GET"])
@login_required
def list_versions():
    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        return jsonify({"versions": []}), 200

    versions = (
        ResumeVersion.query.filter_by(resume_id=resume.id)
        .order_by(ResumeVersion.created_at.desc())
        .all()
    )
    return jsonify(
        {
            "versions": [
                {
                    "id": v.id,
                    "version_name": v.version_name,
                    "target_role": v.target_role or "",
                    "source": v.source or "manual",
                    "ats_score": v.ats_score,
                    "notes": v.notes or "",
                    "created_at": v.created_at.isoformat() if v.created_at else None,
                    "updated_at": v.updated_at.isoformat() if v.updated_at else None,
                }
                for v in versions
            ]
        }
    ), 200


@resume_bp.route("/versions/<int:version_id>", methods=["GET"])
@login_required
def get_version(version_id):
    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        return jsonify({"error": "No resume found"}), 404

    version = ResumeVersion.query.filter_by(id=version_id, resume_id=resume.id).first()
    if not version:
        return jsonify({"error": "Version not found"}), 404

    return jsonify(
        {
            "version": {
                "id": version.id,
                "version_name": version.version_name,
                "target_role": version.target_role or "",
                "source": version.source or "manual",
                "ats_score": version.ats_score,
                "ats_data": version.ats_data,
                "notes": version.notes or "",
                "snapshot": version.snapshot,
                "created_at": version.created_at.isoformat()
                if version.created_at
                else None,
                "updated_at": version.updated_at.isoformat()
                if version.updated_at
                else None,
            }
        }
    ), 200


@resume_bp.route("/versions/<int:version_id>/restore", methods=["POST"])
@login_required
def restore_version(version_id):
    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        return jsonify({"error": "No resume found"}), 404

    version = ResumeVersion.query.filter_by(id=version_id, resume_id=resume.id).first()
    if not version:
        return jsonify({"error": "Version not found"}), 404

    snapshot = version.snapshot
    if not snapshot:
        return jsonify({"error": "Version snapshot is empty"}), 400

    for key, value in snapshot.items():
        if key == "experience":
            value = _normalize_list_field(value, ["bullets", "technologies"])
        elif key == "projects":
            value = _normalize_list_field(value, ["technologies"])
        setattr(resume, key, value)

    db.session.commit()

    from app.resume.profile_bridge import save_resume_to_canonical
    save_resume_to_canonical(current_user.id, snapshot)

    new_version = ResumeVersion(
        resume_id=resume.id,
        version_name=f"{version.version_name} (restored)",
        snapshot=_snapshot(resume),
    )
    db.session.add(new_version)
    db.session.commit()

    from app.core.integration import on_resume_changed
    on_resume_changed(current_user.id)

    from app.career.models import CareerTimelineEvent
    event = CareerTimelineEvent(
        user_id=current_user.id,
        event_type="resume",
        title=f"Resume Restored: {version.version_name}",
        description="Restored a previous version of your resume",
        event_date=datetime.now(timezone.utc),
        importance=2,
    )
    db.session.add(event)
    db.session.commit()

    return jsonify({"message": "Version restored", "resume": _serialize(resume)}), 200


@resume_bp.route("/versions/<int:version_id>", methods=["DELETE"])
@login_required
def delete_version(version_id):
    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        return jsonify({"error": "No resume found"}), 404

    version = ResumeVersion.query.filter_by(id=version_id, resume_id=resume.id).first()
    if not version:
        return jsonify({"error": "Version not found"}), 404

    db.session.delete(version)
    db.session.commit()
    return jsonify({"message": "Version deleted"}), 200


@resume_bp.route("/versions/<int:version_id>", methods=["PUT"])
@login_required
def update_version(version_id):
    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        return jsonify({"error": "No resume found"}), 404
    version = ResumeVersion.query.filter_by(id=version_id, resume_id=resume.id).first()
    if not version:
        return jsonify({"error": "Version not found"}), 404
    data = request.get_json(silent=True) or {}
    if "version_name" in data:
        version.version_name = data["version_name"]
    if "target_role" in data:
        version.target_role = data["target_role"]
    if "notes" in data:
        version.notes = data["notes"]
    if "snapshot" in data:
        snapshot = data["snapshot"]
        if isinstance(snapshot, dict):
            version.snapshot = snapshot
    db.session.commit()
    return jsonify({"message": "Version updated"}), 200


@resume_bp.route("/versions/<int:version_id>/duplicate", methods=["POST"])
@login_required
def duplicate_version(version_id):
    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        return jsonify({"error": "No resume found"}), 404
    source = ResumeVersion.query.filter_by(id=version_id, resume_id=resume.id).first()
    if not source:
        return jsonify({"error": "Version not found"}), 404
    data = request.get_json(silent=True) or {}
    new_name = data.get("version_name", f"{source.version_name} (copy)")
    dup = ResumeVersion(
        resume_id=resume.id,
        version_name=new_name,
        target_role=source.target_role,
        source="manual",
        ats_score=source.ats_score,
        ats_data=source.ats_data,
        snapshot=source.snapshot,
    )
    db.session.add(dup)
    db.session.commit()
    return jsonify({"message": "Version duplicated", "id": dup.id}), 201


@resume_bp.route("/versions/<int:version_id>/tailor", methods=["POST"])
@login_required
@limiter.limit("5 per minute")
def tailor_version(version_id):
    from app.resume.ai_resume_generator import generate_tailored_resume

    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        return jsonify({"error": "No resume found"}), 404
    version = ResumeVersion.query.filter_by(id=version_id, resume_id=resume.id).first()
    if not version:
        return jsonify({"error": "Version not found"}), 404
    data = request.get_json(silent=True) or {}
    job_description = data.get("job_description", "").strip()
    if not job_description:
        return jsonify({"error": "job_description is required"}), 400

    tailored = generate_tailored_resume(current_user.id, version.id, job_description)
    if not tailored:
        return jsonify({"error": "Failed to tailor resume"}), 500

    version_name = data.get("version_name") or f"{version.version_name} (tailored)"
    new_version = ResumeVersion(
        resume_id=resume.id,
        version_name=version_name,
        target_role=version.target_role,
        source="tailored",
        tailored_for_job=job_description[:1000],
        snapshot=tailored,
    )
    db.session.add(new_version)
    db.session.commit()

    from app.core.integration import on_resume_changed
    on_resume_changed(current_user.id)

    return jsonify({
        "message": "Tailored version created",
        "version": {
            "id": new_version.id,
            "version_name": new_version.version_name,
            "source": new_version.source,
        }
    }), 201


@resume_bp.route("/versions/<int:version_id>/ats-score", methods=["POST"])
@login_required
def score_version_ats(version_id):
    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        return jsonify({"error": "No resume found"}), 404
    version = ResumeVersion.query.filter_by(id=version_id, resume_id=resume.id).first()
    if not version:
        return jsonify({"error": "Version not found"}), 404

    data = request.get_json(silent=True) or {}
    job_description = data.get("job_description", "").strip()
    if not job_description:
        return jsonify({"error": "job_description is required"}), 400

    from app.resume.ats import score_resume as ats_score_resume

    snapshot = version.snapshot or {}
    result = ats_score_resume(_snapshot_to_resume_like(snapshot, resume.id, current_user.id), job_description)

    version.ats_score = result.get("overall_score")
    version.ats_data = result
    version.tailored_for_job = job_description[:1000]
    db.session.commit()

    return jsonify(result), 200


def _snapshot_to_resume_like(snapshot, resume_id, user_id):
    """Wrap a snapshot dict in a simple object with attrs that ats.py can read."""
    class ResumeLike:
        pass
    r = ResumeLike()
    r.id = resume_id
    r.user_id = user_id
    r.full_name = snapshot.get("full_name", "")
    r.email = snapshot.get("email", "")
    r.phone = snapshot.get("phone", "")
    r.location = snapshot.get("location", "")
    r.summary = snapshot.get("summary", "")
    r.title = snapshot.get("title", "")
    r.website = snapshot.get("website", "")
    r.linkedin = snapshot.get("linkedin", "")
    r.github = snapshot.get("github", "")
    r.portfolio = snapshot.get("portfolio", "")
    r.experience = snapshot.get("experience", [])
    r.education = snapshot.get("education", [])
    r.projects = snapshot.get("projects", [])
    r.skills = snapshot.get("skills", [])
    r.certificates = snapshot.get("certificates", [])
    r.achievements = snapshot.get("achievements", [])
    r.languages = snapshot.get("languages", [])
    r.publications = snapshot.get("publications", [])
    r.tone = snapshot.get("tone", "professional")
    r.target_job_description = ""
    return r


@resume_bp.route("/generate", methods=["POST"])
@login_required
@limiter.limit("3 per minute")
def generate_ai_resume():
    from app.resume.ai_resume_generator import generate_resume

    data = request.get_json(silent=True) or {}
    target_role = data.get("target_role", "")
    job_description = data.get("job_description", "")

    resume_data = generate_resume(current_user.id, target_role=target_role, job_description=job_description)

    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        resume = Resume(user_id=current_user.id)
        db.session.add(resume)
        db.session.flush()

    for field in ["full_name", "email", "phone", "location", "summary", "title"]:
        if field in resume_data:
            setattr(resume, field, resume_data[field])

    for field in ["experience", "education", "projects", "skills", "certificates", "languages"]:
        if field in resume_data:
            val = resume_data[field]
            if field == "experience":
                val = _normalize_list_field(val, ["bullets", "technologies"])
            elif field == "projects":
                val = _normalize_list_field(val, ["technologies"])
            setattr(resume, field, val)

    db.session.commit()

    from app.resume.profile_bridge import save_resume_to_canonical
    save_resume_to_canonical(current_user.id, resume_data)

    version_name = data.get("version_name") or (f"AI: {target_role}" if target_role else "AI Generated")
    version = ResumeVersion(
        resume_id=resume.id,
        version_name=version_name,
        target_role=target_role or resume_data.get("title", ""),
        source="ai_generated",
        snapshot=_snapshot(resume),
    )
    db.session.add(version)
    db.session.commit()

    from app.core.integration import on_resume_changed
    on_resume_changed(current_user.id)

    return jsonify({
        "message": "Resume generated",
        "version": {
            "id": version.id,
            "version_name": version.version_name,
            "source": version.source,
        }
    }), 201


@resume_bp.route("/versions/compare", methods=["POST"])
@login_required
def compare_versions():
    data = request.get_json(silent=True) or {}
    version_ids = data.get("version_ids", [])

    if len(version_ids) != 2:
        return jsonify({"error": "Provide exactly two version_ids to compare"}), 400

    resume = Resume.query.filter_by(user_id=current_user.id).first()
    if not resume:
        return jsonify({"error": "No resume found"}), 404

    v1 = ResumeVersion.query.filter_by(id=version_ids[0], resume_id=resume.id).first()
    v2 = ResumeVersion.query.filter_by(id=version_ids[1], resume_id=resume.id).first()

    if not v1 or not v2:
        return jsonify({"error": "Version not found"}), 404

    s1, s2 = v1.snapshot or {}, v2.snapshot or {}

    diffs = {}
    for key in set(list(s1.keys()) + list(s2.keys())):
        if s1.get(key) != s2.get(key):
            diffs[key] = {"from": s1.get(key), "to": s2.get(key)}

    return jsonify(
        {
            "version_a": {"id": v1.id, "version_name": v1.version_name},
            "version_b": {"id": v2.id, "version_name": v2.version_name},
            "differences": diffs,
        }
    ), 200


# ── AI Resume Assistant ───────────────────────────────────


@resume_bp.route("/ai/improve-summary", methods=["POST"])
@login_required
@limiter.limit("10 per minute")
def improve_summary():
    from app.ai_service import generate_text, sanitize_for_prompt

    data = request.get_json(silent=True) or {}
    summary = sanitize_for_prompt(data.get("summary", ""))
    tone = sanitize_for_prompt(data.get("tone", "professional"))
    skills = [sanitize_for_prompt(s) for s in data.get("skills", [])]

    prompt = f"""Current summary: "{summary}"
Tone: {tone}
Skills: {", ".join(skills) if skills else "Not specified"}

Rewrite this professional summary to be more impactful and ATS-friendly. Return ONLY the rewritten summary text, no explanations."""

    system = "You are an expert resume writer. Rewrite the professional summary to be stronger, more concise, and optimized for ATS systems."
    result = generate_text(prompt, model="gemini", system_instruction=system)
    return jsonify({"result": result.strip()}), 200


@resume_bp.route("/ai/rewrite-bullet", methods=["POST"])
@login_required
@limiter.limit("10 per minute")
def rewrite_bullet():
    from app.ai_service import generate_text, sanitize_for_prompt

    data = request.get_json(silent=True) or {}
    bullet = sanitize_for_prompt(data.get("bullet", ""))
    tone = sanitize_for_prompt(data.get("tone", "professional"))
    context = sanitize_for_prompt(data.get("context", ""))

    prompt = f"""Original bullet point: "{bullet}"
Role context: {context or "Not specified"}
Tone: {tone}

Rewrite this resume bullet point to be more impactful. Use strong action verbs and include metrics where possible.
Then explain why the rewrite is stronger.

Return in JSON format:
{{
  "original": "...",
  "rewritten": "...",
  "explanation": "..."
}}"""

    system = "You are an expert resume coach. Transform weak bullet points into powerful, metric-driven achievements."
    result = generate_text(prompt, model="gemini", system_instruction=system)

    cleaned = result.strip()
    if cleaned.startswith("```"):
        parts = cleaned.split("```")
        cleaned = parts[1] if len(parts) > 1 else cleaned
        if cleaned.startswith("json"):
            cleaned = cleaned[4:]

    try:
        return jsonify(json.loads(cleaned)), 200
    except json.JSONDecodeError:
        return jsonify(
            {
                "original": bullet,
                "rewritten": result.strip(),
                "explanation": "AI generated response.",
            }
        ), 200


@resume_bp.route("/ai/ats-optimize", methods=["POST"])
@login_required
@limiter.limit("5 per minute")
def ats_optimize():
    from app.ai_service import generate_text, sanitize_for_prompt

    data = request.get_json(silent=True) or {}
    text = sanitize_for_prompt(data.get("text", ""))
    job_description = sanitize_for_prompt(data.get("job_description", ""))

    prompt = f"""Resume text: "{text}"
Target job description: "{job_description}"

Rewrite this resume content to optimize it for ATS matching against the job description.
Incorporate relevant keywords naturally. Preserve the original meaning and truthfulness.

Return in JSON format:
{{
  "original": "...",
  "optimized": "...",
  "keywords_added": ["...", "..."],
  "changes": ["...", "..."]
}}"""

    system = "You are an ATS optimization expert. Rewrite resume content to maximize keyword matching while keeping it natural and truthful."
    result = generate_text(prompt, model="gemini", system_instruction=system)

    cleaned = result.strip()
    if cleaned.startswith("```"):
        parts = cleaned.split("```")
        cleaned = parts[1] if len(parts) > 1 else cleaned
        if cleaned.startswith("json"):
            cleaned = cleaned[4:]

    try:
        return jsonify(json.loads(cleaned)), 200
    except json.JSONDecodeError:
        return jsonify(
            {
                "original": text,
                "optimized": result.strip(),
                "keywords_added": [],
                "changes": [],
            }
        ), 200


@resume_bp.route("/ai/generate-achievement", methods=["POST"])
@login_required
@limiter.limit("10 per minute")
def generate_achievement():
    from app.ai_service import generate_text

    data = request.get_json(silent=True) or {}
    role = data.get("role", "")
    company = data.get("company", "")
    context = data.get("context", "")

    prompt = f"""Role: {role}
Company: {company}
Context: {context or "Not specified"}

Generate 3 measurable achievement bullet points for this role. Each should include a strong action verb and a specific metric.

Return as JSON list:
["...", "...", "..."]"""

    system = "You are an expert resume writer. Generate specific, measurable achievement bullet points."
    result = generate_text(prompt, model="gemini", system_instruction=system)

    cleaned = result.strip()
    if cleaned.startswith("```"):
        parts = cleaned.split("```")
        cleaned = parts[1] if len(parts) > 1 else cleaned
        if cleaned.startswith("json"):
            cleaned = cleaned[4:]

    try:
        return jsonify({"achievements": json.loads(cleaned)}), 200
    except json.JSONDecodeError:
        return jsonify({"achievements": [result.strip()]}), 200


@resume_bp.route("/ai/change-tone", methods=["POST"])
@login_required
@limiter.limit("10 per minute")
def change_tone():
    from app.ai_service import generate_text, sanitize_for_prompt

    data = request.get_json(silent=True) or {}
    text = sanitize_for_prompt(data.get("text", ""))
    target_tone = sanitize_for_prompt(data.get("tone", "professional"))

    prompt = f"""Text: "{text}"
Target tone: {target_tone}

Rewrite this resume content to have a {target_tone} tone while keeping the same information.

Return the rewritten text only."""

    system = f"You are an expert writer. Rewrite the given resume text in a {target_tone} tone. Only return the rewritten text."
    result = generate_text(prompt, model="gemini", system_instruction=system)
    return jsonify({"result": result.strip()}), 200


# ── Cover Letter Generator ────────────────────────────────


@resume_bp.route("/cover-letter", methods=["POST"])
@login_required
@limiter.limit("5 per minute")
def generate_cover_letter():
    from app.ai_service import generate_text, sanitize_for_prompt

    data = request.get_json(silent=True) or {}
    company = sanitize_for_prompt(data.get("company", ""))
    role = sanitize_for_prompt(data.get("role", ""))
    job_description = sanitize_for_prompt(data.get("job_description", ""))
    tone = sanitize_for_prompt(data.get("tone", "professional"))

    resume = Resume.query.filter_by(user_id=current_user.id).first()
    resume_text = ""
    if resume:
        skills = ", ".join(resume.skills or [])
        summary = resume.summary or ""
        exp_items = []
        for exp in resume.experience or []:
            bullets = exp.get("bullets", "")
            if isinstance(bullets, list):
                bullets = "; ".join(bullets)
            exp_items.append(
                f"{exp.get('role', '')} at {exp.get('company', '')}: {bullets}"
            )
        experience = "\n".join(exp_items)
        resume_text = f"Summary: {summary}\nSkills: {skills}\nExperience:\n{experience}"

    prompt = f"""Generate a professional cover letter.

Company: {company}
Role: {role}
Job Description: {job_description or "Not provided"}
Tone: {tone}

Applicant Background:
{resume_text or "No resume data available"}

Write a compelling cover letter that highlights relevant experience and skills."""

    system = f"You are an expert cover letter writer. Write a {tone} cover letter that connects the applicant's experience to the role."
    result = generate_text(prompt, model="gemini", system_instruction=system)
    return jsonify({"cover_letter": result.strip()}), 200
